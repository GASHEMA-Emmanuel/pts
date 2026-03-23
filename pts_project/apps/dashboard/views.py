"""
Dashboard views for template-based web interface.
"""
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.db.models import Count, Sum, Q
from django.utils import timezone
from django.conf import settings
from datetime import timedelta
from functools import wraps
from docx import Document as WordDocument
from openpyxl import load_workbook
import xlrd
import html
import io


def convert_word_to_html(file_path):
    """
    Convert a Word (.docx) document to HTML.
    Returns HTML string or error message.
    """
    try:
        doc = WordDocument(file_path)
        html_content = '<div style="font-family: Arial, sans-serif; line-height: 1.6; max-width: 100%;">'
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Get paragraph style level for indentation
                level = para.style.name.split()[-1] if 'Heading' in para.style.name else '0'
                html_content += f'<p style="margin: 10px 0;">{html.escape(para.text)}</p>'
        
        # Add tables
        for table in doc.tables:
            html_content += '<table style="border-collapse: collapse; width: 100%; margin: 15px 0;">'
            for row in table.rows:
                html_content += '<tr>'
                for cell in row.cells:
                    html_content += f'<td style="border: 1px solid #ddd; padding: 10px;">{html.escape(cell.text)}</td>'
                html_content += '</tr>'
            html_content += '</table>'
        
        html_content += '</div>'
        return html_content
    except Exception as e:
        return f'<div class="alert alert-danger">Error reading Word document: {str(e)}</div>'


def convert_excel_to_html(file_path):
    """
    Convert an Excel (.xlsx/.xls) file to HTML tables.
    Supports both old .xls (xlrd) and new .xlsx (openpyxl) formats.
    Falls back to CSV reading if file format is misidentified.
    Returns HTML string or error message.
    """
    import csv
    
    try:
        html_content = '<div style="font-family: Arial, sans-serif; max-width: 100%; overflow-x: auto;">'
        success = False
        
        # Check file format and use appropriate library
        if file_path.lower().endswith('.xls'):
            # Old Excel format - try xlrd first
            try:
                wb = xlrd.open_workbook(file_path)
                for sheet_index, sheet_name in enumerate(wb.sheet_names()):
                    ws = wb.sheet_by_index(sheet_index)
                    html_content += f'<h3 style="margin-top: 20px; margin-bottom: 10px; color: #004085;">{html.escape(sheet_name)}</h3>'
                    html_content += '<table style="border-collapse: collapse; margin-bottom: 20px; border: 1px solid #ddd;">'
                    
                    for row_idx in range(ws.nrows):
                        html_content += '<tr>'
                        for col_idx in range(ws.ncols):
                            cell = ws.cell(row_idx, col_idx)
                            cell_text = str(cell.value) if cell.value is not None else ''
                            html_content += f'<td style="border: 1px solid #ddd; padding: 8px; background-color: #f8f9fa;">{html.escape(cell_text)}</td>'
                        html_content += '</tr>'
                    
                    html_content += '</table>'
                success = True
            except Exception as xlrd_error:
                # xlrd failed, try openpyxl as fallback
                try:
                    wb = load_workbook(file_path)
                    for sheet_name in wb.sheetnames:
                        ws = wb[sheet_name]
                        html_content += f'<h3 style="margin-top: 20px; margin-bottom: 10px; color: #004085;">{html.escape(sheet_name)}</h3>'
                        html_content += '<table style="border-collapse: collapse; margin-bottom: 20px; border: 1px solid #ddd;">'
                        
                        for row in ws.iter_rows(values_only=True):
                            html_content += '<tr>'
                            for cell_value in row:
                                cell_text = str(cell_value) if cell_value is not None else ''
                                html_content += f'<td style="border: 1px solid #ddd; padding: 8px; background-color: #f8f9fa;">{html.escape(cell_text)}</td>'
                            html_content += '</tr>'
                        
                        html_content += '</table>'
                    success = True
                except Exception:
                    # Both xlrd and openpyxl failed, try CSV as last resort
                    pass
        else:
            # New Excel format (.xlsx) - try openpyxl first
            try:
                wb = load_workbook(file_path)
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    html_content += f'<h3 style="margin-top: 20px; margin-bottom: 10px; color: #004085;">{html.escape(sheet_name)}</h3>'
                    html_content += '<table style="border-collapse: collapse; margin-bottom: 20px; border: 1px solid #ddd;">'
                    
                    for row in ws.iter_rows(values_only=True):
                        html_content += '<tr>'
                        for cell_value in row:
                            cell_text = str(cell_value) if cell_value is not None else ''
                            html_content += f'<td style="border: 1px solid #ddd; padding: 8px; background-color: #f8f9fa;">{html.escape(cell_text)}</td>'
                        html_content += '</tr>'
                    
                    html_content += '</table>'
                success = True
            except Exception:
                # openpyxl failed, try xlrd as fallback
                try:
                    wb = xlrd.open_workbook(file_path)
                    for sheet_index, sheet_name in enumerate(wb.sheet_names()):
                        ws = wb.sheet_by_index(sheet_index)
                        html_content += f'<h3 style="margin-top: 20px; margin-bottom: 10px; color: #004085;">{html.escape(sheet_name)}</h3>'
                        html_content += '<table style="border-collapse: collapse; margin-bottom: 20px; border: 1px solid #ddd;">'
                        
                        for row_idx in range(ws.nrows):
                            html_content += '<tr>'
                            for col_idx in range(ws.ncols):
                                cell = ws.cell(row_idx, col_idx)
                                cell_text = str(cell.value) if cell.value is not None else ''
                                html_content += f'<td style="border: 1px solid #ddd; padding: 8px; background-color: #f8f9fa;">{html.escape(cell_text)}</td>'
                            html_content += '</tr>'
                        
                        html_content += '</table>'
                    success = True
                except Exception:
                    # Both formats failed
                    pass
        
        # Final fallback: try reading as CSV
        if not success:
            try:
                html_content += '<h3 style="margin-top: 20px; margin-bottom: 10px; color: #004085;">Data (CSV Format)</h3>'
                html_content += '<table style="border-collapse: collapse; margin-bottom: 20px; border: 1px solid #ddd;">'
                
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as csvfile:
                    csv_reader = csv.reader(csvfile)
                    for row in csv_reader:
                        html_content += '<tr>'
                        for cell_value in row:
                            html_content += f'<td style="border: 1px solid #ddd; padding: 8px; background-color: #f8f9fa;">{html.escape(cell_value)}</td>'
                        html_content += '</tr>'
                
                html_content += '</table>'
                success = True
            except Exception:
                pass
        
        if not success:
            return '<div class="alert alert-danger">Unable to preview this file. Please download it to view the contents.</div>'
        
        html_content += '</div>'
        return html_content
    except Exception as e:
        return f'<div class="alert alert-danger">Error reading file: {str(e)}</div>'


@login_required
def dashboard_view(request):
    """
    Main dashboard - routes to role-specific dashboard.
    Admin users see admin dashboard.
    CBM users see CBM dashboard.
    HOD/DM users see HOD dashboard.
    Procurement Team users see procurement dashboard.
    Other users see their division's submissions.
    """
    user = request.user
    
    # Route based on user role
    if user.role and user.role.name == 'Admin':
        return admin_dashboard_view(request)
    elif user.role and user.role.name == 'CBM':
        return cbm_dashboard_view(request)
    elif user.role and user.role.name == 'HOD/DM':
        return hod_dashboard_view(request)
    elif user.role and user.role.name == 'Procurement Team':
        # Import here to avoid circular imports
        from apps.dashboard.views_procurement import procurement_dashboard_view
        return procurement_dashboard_view(request)
    else:
        # Show user's division submissions
        return division_dashboard_view(request)


@login_required
def division_dashboard_view(request):
    """
    Division user dashboard showing their submissions.
    """
    user = request.user
    
    # Import models here to avoid circular imports
    from apps.procurement.models import ProcurementCall, Submission, Bid
    from apps.workflows.models import WorkflowStage, WorkflowHistory
    from apps.notifications.models import Notification
    
    # Get base querysets
    calls = ProcurementCall.objects.all()
    submissions = Submission.objects.filter(is_deleted=False)
    
    # Filter by division for non-admin users
    if user.division:
        submissions = submissions.filter(division=user.division)
    
    # Calculate KPI stats
    stats = {
        'active_calls': calls.filter(status='Active').count(),
        'pending_approvals': submissions.filter(
            status__in=['HOD/DM Submit', 'Review of Procurement Draft', 'CBM Review', 'Prepare Tender Document', 'CBM Review TD']
        ).count(),
        'total_submissions': submissions.count(),
        'pending_bids': Bid.objects.filter(
            submission__status='Bidding',
            is_winner=False,
            is_disqualified=False
        ).count(),
    }
    
    # Get pending approvals for table
    pending_approvals = []
    pending_subs = submissions.filter(
        status__in=['HOD/DM Submit', 'Review of Procurement Draft', 'CBM Review', 'Prepare Tender Document', 'CBM Review TD']
    ).select_related('current_stage')[:5]
    
    for sub in pending_subs:
        stage_class = 'review'
        if sub.current_stage:
            if 'approval' in sub.current_stage.name.lower():
                stage_class = 'approval'
            elif 'finance' in sub.current_stage.name.lower():
                stage_class = 'finance'
        
        pending_approvals.append({
            'tracking_reference': sub.tracking_reference,
            'current_stage_name': sub.current_stage.name if sub.current_stage else 'Unknown',
            'stage_class': stage_class,
            'days_at_stage': sub.days_at_current_stage,
        })
    
    # Get recent submissions
    recent_submissions = []
    recent_subs = submissions.order_by('-created_at')[:5]
    
    for sub in recent_subs:
        status_class = 'pending'
        if sub.status == 'Publish Plan':
            status_class = 'approved'
        elif sub.status in ['Review of Procurement Draft', 'CBM Review', 'Prepare Tender Document', 'CBM Review TD']:
            status_class = 'under-review'
        elif sub.status == 'Rejected':
            status_class = 'danger'
        
        recent_submissions.append({
            'item_name': sub.item_name,
            'tracking_reference': sub.tracking_reference,
            'status': sub.status,
            'status_class': status_class,
        })
    
    # Get recent activity from workflow history
    recent_activity = []
    histories = WorkflowHistory.objects.select_related(
        'submission', 'action_by'
    ).order_by('-created_at')[:5]
    
    for history in histories:
        activity_type = 'info'
        title = history.get_action_display()
        
        if history.action == 'approve':
            activity_type = 'success'
            title = 'Submission Plan Published'
        elif history.action == 'reject':
            activity_type = 'danger'
            title = 'Submission Rejected'
        elif history.action == 'submit':
            activity_type = 'info'
            title = 'New Submission Created'
        elif history.action == 'return':
            activity_type = 'warning'
            title = 'Approval Requested'
        
        recent_activity.append({
            'type': activity_type,
            'title': title,
            'description': f"{history.submission.tracking_reference} by {history.action_by.full_name if history.action_by else 'System'}",
            'time': time_ago(history.created_at),
        })
    
    # Get workflow summary
    workflow_summary = []
    for stage in WorkflowStage.objects.all():
        count = submissions.filter(current_stage=stage).count()
        workflow_summary.append({
            'name': stage.name,
            'count': count,
            'color': stage.color,
        })
    
    # Get unread notifications count
    unread_notifications_count = Notification.objects.filter(
        user=user,
        is_read=False
    ).count()
    
    context = {
        'stats': stats,
        'pending_approvals': pending_approvals,
        'recent_submissions': recent_submissions,
        'recent_activity': recent_activity,
        'workflow_summary': workflow_summary,
        'unread_notifications_count': unread_notifications_count,
    }
    
    return render(request, 'dashboard/index.html', context)


def time_ago(dt):
    """Calculate human-readable time difference."""
    now = timezone.now()
    diff = now - dt
    
    if diff.days > 0:
        if diff.days == 1:
            return '1 day ago'
        return f'{diff.days} days ago'
    
    hours = diff.seconds // 3600
    if hours > 0:
        if hours == 1:
            return '1 hour ago'
        return f'{hours} hours ago'
    
    minutes = diff.seconds // 60
    if minutes > 0:
        if minutes == 1:
            return '1 minute ago'
        return f'{minutes} minutes ago'
    
    return 'Just now'


def admin_required(view_func):
    """Decorator to check if user is an admin."""
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated or (request.user.role and request.user.role.name != 'Admin'):
            return redirect('dashboard')
        return view_func(request, *args, **kwargs)
    return wrapper


@login_required
@admin_required
def admin_dashboard_view(request):
    """
    Admin dashboard showing platform overview and statistics.
    """
    from apps.accounts.models import User, Role, UserActivity
    from apps.divisions.models import Division
    from apps.procurement.models import ProcurementCall, Submission

    # Calculate admin stats
    total_users_active = User.objects.filter(is_active=True).count()
    one_week_ago = timezone.now() - timedelta(days=7)
    active_users_week = User.objects.filter(last_login__gte=one_week_ago, is_active=True).count()

    pending_approvals = Submission.objects.filter(
        is_deleted=False,
        status__in=['HOD/DM Submit', 'Review of Procurement Draft', 'CBM Review', 'Prepare Tender Document', 'CBM Review TD']
    ).count()

    stats = {
        'total_users': total_users_active,
        'total_inactive': User.objects.filter(is_active=False).count(),
        'total_all': User.objects.count(),
        'total_divisions': Division.objects.count(),
        'active_calls': ProcurementCall.objects.filter(status='Active').count(),
        'total_submissions': Submission.objects.filter(is_deleted=False).count(),
        'active_users_week': active_users_week,
        'pending_approvals': pending_approvals,
        'new_users_30days': User.objects.filter(
            created_at__gte=timezone.now() - timedelta(days=30)
        ).count(),
    }

    # User breakdown by role with percentage
    user_by_role = []
    total_all = stats['total_all'] or 1  # avoid division by zero
    for role in Role.objects.all():
        count = User.objects.filter(role=role, is_active=True).count()
        user_by_role.append({
            'name': role.name,
            'count': count,
            'percent': round(count / total_all * 100),
        })

    # Recent user activities
    recent_activities = UserActivity.objects.select_related('user').order_by('-created_at')[:12]

    # Recent new users (last 5)
    recent_users = User.objects.select_related('role', 'division').order_by('-created_at')[:5]

    context = {
        'stats': stats,
        'user_by_role': user_by_role,
        'active_users_week': active_users_week,
        'recent_activities': recent_activities,
        'recent_users': recent_users,
    }

    return render(request, 'admin/dashboard.html', context)


@login_required
@admin_required
def admin_users_view(request):
    """
    Admin user management view.
    Handles GET (list/filter) and POST (create, edit, toggle_active).
    """
    from apps.accounts.models import User, Role
    from django.contrib import messages

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'create_user':
            full_name = request.POST.get('full_name', '').strip()
            email = request.POST.get('email', '').strip()
            password = request.POST.get('password', '')
            password_confirm = request.POST.get('password_confirm', '')
            role_id = request.POST.get('role') or None
            division_id = request.POST.get('division') or None
            phone_number = request.POST.get('phone_number', '').strip()
            is_active = request.POST.get('is_active') == 'on'

            if not full_name or not email or not password:
                messages.error(request, 'Full name, email and password are required.')
            elif password != password_confirm:
                messages.error(request, 'Passwords do not match.')
            elif User.objects.filter(email=email).exists():
                messages.error(request, f'A user with email {email} already exists.')
            else:
                try:
                    user = User.objects.create_user(
                        email=email,
                        password=password,
                        full_name=full_name,
                        phone_number=phone_number or None,
                        is_active=is_active,
                    )
                    if role_id:
                        user.role_id = role_id
                    if division_id:
                        user.division_id = division_id
                    user.save()
                    messages.success(request, f'User "{full_name}" created successfully.')
                except Exception as e:
                    messages.error(request, f'Error creating user: {e}')

        elif action == 'edit_user':
            user_id = request.POST.get('user_id')
            try:
                user = User.objects.get(id=user_id)
                user.full_name = request.POST.get('full_name', user.full_name).strip()
                user.email = request.POST.get('email', user.email).strip()
                user.phone_number = request.POST.get('phone_number', '').strip() or None
                role_id = request.POST.get('role') or None
                division_id = request.POST.get('division') or None
                user.role_id = role_id
                user.division_id = division_id
                # Optionally update password if provided
                new_password = request.POST.get('new_password', '').strip()
                if new_password:
                    user.set_password(new_password)
                user.save()
                messages.success(request, f'User "{user.full_name}" updated successfully.')
            except User.DoesNotExist:
                messages.error(request, 'User not found.')
            except Exception as e:
                messages.error(request, f'Error updating user: {e}')

        elif action == 'toggle_user':
            user_id = request.POST.get('user_id')
            try:
                user = User.objects.get(id=user_id)
                user.is_active = not user.is_active
                user.save(update_fields=['is_active'])
                state = 'activated' if user.is_active else 'deactivated'
                messages.success(request, f'User "{user.full_name}" {state}.')
            except User.DoesNotExist:
                messages.error(request, 'User not found.')

        return redirect('admin_users')

    # GET — list with filters
    users = User.objects.all().select_related('role', 'division').order_by('-created_at')

    # Search by name or email
    search_q = request.GET.get('q', '').strip()
    if search_q:
        users = users.filter(
            Q(full_name__icontains=search_q) | Q(email__icontains=search_q)
        )

    # Filter by role
    role_filter = request.GET.get('role')
    if role_filter:
        users = users.filter(role_id=role_filter)

    # Filter by division
    division_filter = request.GET.get('division')
    if division_filter:
        users = users.filter(division_id=division_filter)

    # Filter by status
    status_filter = request.GET.get('status', 'active')
    if status_filter == 'active':
        users = users.filter(is_active=True)
    elif status_filter == 'inactive':
        users = users.filter(is_active=False)

    roles = Role.objects.all()
    from apps.divisions.models import Division
    divisions = Division.objects.all()

    context = {
        'users': users,
        'roles': roles,
        'divisions': divisions,
        'active_role_filter': role_filter,
        'active_division_filter': division_filter,
        'active_status_filter': status_filter,
        'search_q': search_q,
        'total_all_users': User.objects.count(),
        'total_active': User.objects.filter(is_active=True).count(),
        'total_inactive': User.objects.filter(is_active=False).count(),
    }

    return render(request, 'admin/users.html', context)


@login_required
@admin_required
def admin_divisions_view(request):
    """
    Admin division management view.
    Handles GET (list) and POST (create, edit, delete).
    """
    from apps.divisions.models import Division
    from apps.accounts.models import User
    from django.contrib import messages

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'create_division':
            name = request.POST.get('name', '').strip()
            description = request.POST.get('description', '').strip()
            if not name:
                messages.error(request, 'Division name is required.')
            elif Division.objects.filter(name__iexact=name).exists():
                messages.error(request, f'A division named "{name}" already exists.')
            else:
                try:
                    Division.objects.create(name=name, description=description)
                    messages.success(request, f'Division "{name}" created successfully.')
                except Exception as e:
                    messages.error(request, f'Error creating division: {e}')

        elif action == 'edit_division':
            div_id = request.POST.get('division_id')
            try:
                division = Division.objects.get(id=div_id)
                division.name = request.POST.get('name', division.name).strip()
                division.description = request.POST.get('description', '').strip()
                division.save()
                messages.success(request, f'Division "{division.name}" updated.')
            except Division.DoesNotExist:
                messages.error(request, 'Division not found.')
            except Exception as e:
                messages.error(request, f'Error updating division: {e}')

        elif action == 'delete_division':
            div_id = request.POST.get('division_id')
            try:
                division = Division.objects.get(id=div_id)
                if User.objects.filter(division=division).exists():
                    messages.error(request, f'Cannot delete "{division.name}" — it still has users assigned to it.')
                else:
                    name = division.name
                    division.delete()
                    messages.success(request, f'Division "{name}" deleted.')
            except Division.DoesNotExist:
                messages.error(request, 'Division not found.')

        return redirect('admin_divisions')

    # GET
    divisions = Division.objects.all()
    divisions_data = []
    for div in divisions:
        user_count = User.objects.filter(division=div, is_active=True).count()
        divisions_data.append({
            'id': div.id,
            'name': div.name,
            'description': div.description,
            'user_count': user_count,
            'created_at': div.created_at,
        })

    context = {
        'divisions': divisions_data,
        'total_divisions': len(divisions_data),
    }

    return render(request, 'admin/divisions.html', context)


@login_required
@admin_required
def admin_audit_logs_view(request):
    """
    Admin audit logs view with filtering, date range and pagination.
    """
    from apps.accounts.models import UserActivity, User
    from django.core.paginator import Paginator

    logs = UserActivity.objects.select_related('user').order_by('-created_at')

    # Filter by action
    action_filter = request.GET.get('action')
    if action_filter:
        logs = logs.filter(action=action_filter)

    # Filter by user
    user_filter = request.GET.get('user')
    if user_filter:
        logs = logs.filter(user_id=user_filter)

    # Date range filters
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    if date_from:
        try:
            from datetime import datetime
            logs = logs.filter(created_at__date__gte=datetime.strptime(date_from, '%Y-%m-%d').date())
        except ValueError:
            pass
    if date_to:
        try:
            from datetime import datetime
            logs = logs.filter(created_at__date__lte=datetime.strptime(date_to, '%Y-%m-%d').date())
        except ValueError:
            pass

    total_logs = logs.count()

    # Pagination — 50 per page
    paginator = Paginator(logs, 50)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)

    actions = UserActivity.objects.values_list('action', flat=True).distinct()
    users = User.objects.filter(is_active=True).order_by('full_name')

    context = {
        'logs': page_obj,
        'page_obj': page_obj,
        'total_logs': total_logs,
        'actions': actions,
        'users': users,
        'active_action_filter': action_filter,
        'active_user_filter': user_filter,
        'date_from': date_from or '',
        'date_to': date_to or '',
    }

    return render(request, 'admin/audit_logs.html', context)


@login_required
@admin_required
def admin_system_settings_view(request):
    """
    Admin system settings view.
    """
    # Placeholder for system settings
    # In production, this would load settings from database
    
    settings_sections = [
        {
            'title': 'Notification Thresholds',
            'icon': 'bi-bell',
            'description': 'Configure notification settings and alerts',
            'items': [
                {'name': 'Email Notifications', 'value': 'Enabled'},
                {'name': 'SMS Alerts', 'value': 'Disabled'},
                {'name': 'Approval Escalation', 'value': 'After 3 days'},
            ]
        },
        {
            'title': 'Security Settings',
            'icon': 'bi-shield-lock',
            'description': 'Manage security policies and access controls',
            'items': [
                {'name': 'Password Policy', 'value': 'Strong (12+ chars)'},
                {'name': 'Session Timeout', 'value': '30 minutes'},
                {'name': 'IP Whitelisting', 'value': 'Disabled'},
            ]
        },
        {
            'title': 'System Maintenance',
            'icon': 'bi-gear',
            'description': 'System health and maintenance tasks',
            'items': [
                {'name': 'Last Backup', 'value': 'Today at 2:00 AM'},
                {'name': 'Database Size', 'value': '250 MB'},
                {'name': 'Cache Status', 'value': 'Healthy'},
            ]
        },
    ]
    
    context = {
        'settings_sections': settings_sections,
    }
    
    return render(request, 'admin/system_settings.html', context)


@login_required
@admin_required
def admin_reports_view(request):
    """
    Admin reports view.
    """
    from apps.accounts.models import User
    from apps.procurement.models import ProcurementCall, Submission
    from apps.contracts.models import Contract
    from django.db.models import Q

    # User Statistics
    total_users = User.objects.count()
    active_users_month = User.objects.filter(
        last_login__gte=timezone.now() - timedelta(days=30)
    ).count()
    new_users_30days = User.objects.filter(
        created_at__gte=timezone.now() - timedelta(days=30)
    ).count()
    inactive_users = User.objects.filter(is_active=False).count()

    user_stats = {
        'total_users': total_users,
        'active_this_month': active_users_month,
        'new_users_30days': new_users_30days,
        'inactive_users': inactive_users,
    }

    # Procurement Statistics
    active_calls = ProcurementCall.objects.filter(status='Active').count()
    total_calls = ProcurementCall.objects.count()
    total_submissions = Submission.objects.filter(is_deleted=False).count()
    approved_submissions = Submission.objects.filter(
        is_deleted=False,
        status__in=[
            'Publish Plan', 'Prepare Tender Document', 'CBM Review TD',
            'Publication of TD', 'Opening', 'Evaluation', 'CBM Approval',
            'Notify Bidders', 'Contract Negotiation', 'Contract Drafting',
            'Legal Review', 'Supplier Approval', 'MINIJUST Legal Review',
            'Awarded', 'Completed'
        ]
    ).count()
    pending_approvals = Submission.objects.filter(
        is_deleted=False,
        status__in=['HOD/DM Submit', 'Review of Procurement Draft', 'CBM Review',
                    'Prepare Tender Document', 'CBM Review TD']
    ).count()

    procurement_stats = {
        'total_calls': total_calls,
        'active_calls': active_calls,
        'total_submissions': total_submissions,
        'approved_submissions': approved_submissions,
        'pending_approvals': pending_approvals,
    }

    # Contract statistics
    try:
        total_contracts = Contract.objects.count()
        active_contracts = Contract.objects.filter(status='Active').count()
        completed_contracts = Contract.objects.filter(status='Completed').count()
        cancelled_contracts = Contract.objects.filter(status='Cancelled').count()
        renewed_contracts = Contract.objects.filter(status='Renewed').count()
    except Exception:
        total_contracts = active_contracts = completed_contracts = cancelled_contracts = renewed_contracts = 0

    contract_stats = {
        'total_contracts': total_contracts,
        'active_contracts': active_contracts,
        'completed_contracts': completed_contracts,
        'expired_contracts': cancelled_contracts,
    }

    # Submission status breakdown
    submission_by_status = []
    STATUS_GROUPS = {
        'Draft / Pending HOD': ['HOD/DM Submit'],
        'Under Review': ['Review of Procurement Draft', 'CBM Review'],
        'Tender Stage': ['Prepare Tender Document', 'CBM Review TD', 'Publication of TD', 'Opening'],
        'Evaluation': ['Evaluation', 'CBM Approval'],
        'Contract Stage': ['Contract Negotiation', 'Contract Drafting', 'Legal Review',
                           'Supplier Approval', 'MINIJUST Legal Review'],
        'Awarded / Completed': ['Awarded', 'Completed'],
    }
    for group_name, statuses in STATUS_GROUPS.items():
        count = Submission.objects.filter(is_deleted=False, status__in=statuses).count()
        submission_by_status.append({'name': group_name, 'count': count})

    context = {
        'user_stats': user_stats,
        'procurement_stats': procurement_stats,
        'contract_stats': contract_stats,
        'submission_by_status': submission_by_status,
    }

    return render(request, 'admin/reports.html', context)


# CBM Dashboard Views
@login_required
def cbm_dashboard_view(request):
    """
    CBM dashboard showing procurement oversight and strategic management.
    """
    user = request.user

    # Redirect non-CBM users
    if user.role and user.role.name != 'CBM':
        return redirect('dashboard')

    from apps.procurement.models import ProcurementCall, Submission, Tender as _Tender
    from apps.contracts.models import Contract
    from apps.divisions.models import Division
    from apps.workflows.models import WorkflowHistory
    from apps.procurement.models import CompiledDocument
    from django.db.models import Sum, Count

    # ── Active calls ─────────────────────────────────────────────
    active_calls_qs = ProcurementCall.objects.filter(status='Active')

    # ── Stages where CBM is notified (auto-advanced, no manual action required) ──
    # Note: These stages now auto-advance with notification to CBM for information only
    CBM_ACTION_STAGES = ['CBM Review', 'CBM Review TD', 'CBM Approval']

    # ── Submissions awaiting CBM action ──────────────────────────
    # Note: This queue will be empty as submissions auto-advance through CBM stages
    action_queue_qs = Submission.objects.filter(
        status__in=CBM_ACTION_STAGES,
        is_deleted=False
    ).select_related('division', 'call').order_by('-updated_at')

    action_queue = []
    for sub in action_queue_qs:
        action_queue.append({
            'id': sub.id,
            'tracking_reference': sub.tracking_reference,
            'item_name': sub.item_name,
            'division': sub.division.name if sub.division else '—',
            'status': sub.status,
            'total_budget': sub.total_budget,
            'days_waiting': sub.days_at_current_stage,
            'call_ref': sub.call.reference_number if sub.call else '—',
        })

    # ── Contract stats ───────────────────────────────────────────
    active_contracts_qs = Contract.objects.filter(status__in=['Active', 'Renewed'])
    overdue_contracts = []
    overdue_count = 0
    contract_health = []

    for c in active_contracts_qs.select_related('division'):
        prog = c.lumpsum_progress_data
        days = c.days_until_delivery
        is_overdue = (prog and prog.get('is_overdue')) or (days is not None and days < 0)
        is_warning = (prog and prog.get('is_quarter_alert')) or (days is not None and 0 <= days <= 30)

        if is_overdue:
            overdue_count += 1
            health = 'overdue'
        elif is_warning:
            health = 'warning'
        else:
            health = 'healthy'

        contract_health.append({
            'id': c.id,
            'contract_number': c.contract_number,
            'contract_name': c.contract_name,
            'contract_type': c.contract_type,
            'status': c.status,
            'health': health,
            'days_until_delivery': days,
            'division': c.division.name if c.division else '—',
        })

    # sort: overdue first, then warning, then healthy
    _order = {'overdue': 0, 'warning': 1, 'healthy': 2}
    contract_health.sort(key=lambda x: _order[x['health']])

    # ── Pipeline funnel ──────────────────────────────────────────
    ALL_PIPELINE_STATUSES = [
        'HOD/DM Submit', 'Review of Procurement Draft', 'Submit Compiled Document',
        'CBM Review', 'Publish Plan', 'Prepare Tender Document', 'CBM Review TD',
        'Publication of TD', 'Opening', 'Evaluation', 'CBM Approval',
        'Notify Bidders', 'Contract Negotiation', 'Contract Drafting',
        'Legal Review', 'Supplier Approval', 'MINIJUST Legal Review',
        'Awarded', 'Completed',
    ]
    pipeline_funnel = []
    for status in ALL_PIPELINE_STATUSES:
        count = Submission.objects.filter(status=status, is_deleted=False).count()
        pipeline_funnel.append({'status': status, 'count': count})

    # ── Total budget under management ───────────────────────────
    live_statuses = [s for s in ALL_PIPELINE_STATUSES if s not in ('Completed', 'Awarded')]
    total_budget = Submission.objects.filter(
        status__in=live_statuses, is_deleted=False
    ).aggregate(total=Sum('total_budget'))['total'] or 0

    # ── Active calls with per-division progress ──────────────────
    active_calls_data = []
    for call in active_calls_qs.prefetch_related('submissions'):
        subs = Submission.objects.filter(call=call, is_deleted=False)
        active_calls_data.append({
            'id': call.id,
            'reference_number': call.reference_number,
            'title': call.title,
            'end_date': call.end_date,
            'days_remaining': call.days_remaining,
            'total_submissions': subs.count(),
            'divisions_submitted': subs.values('division').distinct().count(),
        })

    # ── Recent workflow activity ─────────────────────────────────
    recent_activity = []
    for h in WorkflowHistory.objects.select_related(
        'submission', 'action_by', 'from_stage', 'to_stage'
    ).order_by('-created_at')[:10]:
        recent_activity.append({
            'actor': h.action_by.full_name if h.action_by else 'System',
            'action': h.get_action_display(),
            'submission_ref': h.submission.tracking_reference if h.submission else '—',
            'from_stage': h.from_stage.name if h.from_stage else '—',
            'to_stage': h.to_stage.name if h.to_stage else '—',
            'time': time_ago(h.created_at),
        })

    # ── Division submission snapshot ─────────────────────────────
    divisions_status = []
    for division in Division.objects.all():
        subs = Submission.objects.filter(division=division, is_deleted=False)
        divisions_status.append({
            'name': division.name,
            'submission_count': subs.count(),
            'cbm_action': subs.filter(status__in=CBM_ACTION_STAGES).count(),
            'awarded': subs.filter(status__in=['Awarded', 'Completed']).count(),
        })

    # ── KPI stats ────────────────────────────────────────────────
    stats = {
        'active_calls': active_calls_qs.count(),
        'awaiting_cbm_action': len(action_queue),
        'tenders_in_progress': _Tender.objects.exclude(
            status__in=['Completed', 'Cancelled']
        ).count(),
        'active_contracts': active_contracts_qs.count(),
        'overdue_contracts': overdue_count,
        'total_budget': total_budget,
        'total_divisions': Division.objects.count(),
    }

    context = {
        'stats': stats,
        'action_queue': action_queue,
        'pipeline_funnel': pipeline_funnel,
        'contract_health': contract_health,
        'active_calls_data': active_calls_data,
        'divisions_status': divisions_status,
        'recent_activity': recent_activity,
        'pending_compiled_count': CompiledDocument.objects.filter(status='Sent to CBM').count(),
    }

    return render(request, 'cbm/dashboard.html', context)


@login_required
def cbm_calls_view(request):
    """
    CBM procurement calls management view.
    """
    user = request.user
    
    if user.role and user.role.name != 'CBM':
        return redirect('dashboard')
    
    from apps.procurement.models import ProcurementCall, Submission
    
    calls = ProcurementCall.objects.all().order_by('-created_at')
    
    # Add submission stats to each call
    calls_data = []
    for call in calls:
        submissions = Submission.objects.filter(
            call=call,
            is_deleted=False
        )
        calls_data.append({
            'id': call.id,
            'title': call.title,
            'description': call.description,
            'status': call.status,
            'deadline': call.end_date,
            'created_at': call.created_at,
            'total_submissions': submissions.count(),
            'pending_submissions': submissions.filter(
                status__in=['HOD/DM Submit', 'Review of Procurement Draft', 'CBM Review']
            ).count(),
        })
    
    context = {
        'calls': calls_data,
    }
    
    return render(request, 'cbm/calls.html', context)


@login_required
def cbm_submissions_view(request):
    """
    CBM submissions review view.
    Shows ALL submissions from ALL divisions regardless of status.
    CBM has visibility into all submissions regardless of division or stage.
    """
    user = request.user
    
    if user.role and user.role.name != 'CBM':
        return redirect('dashboard')
    
    from apps.procurement.models import Submission
    
    # Get ALL submissions from ALL divisions - CBM has full visibility
    submissions = Submission.objects.filter(
        is_deleted=False
    ).select_related('division', 'current_stage').order_by('-created_at')
    
    # Filter by status if specified (optional)
    status_filter = request.GET.get('status')
    if status_filter:
        submissions = submissions.filter(status=status_filter)
    
    # Filter by division if specified (optional - for user convenience)
    division_filter = request.GET.get('division')
    if division_filter:
        submissions = submissions.filter(division_id=division_filter)
    
    from apps.divisions.models import Division
    divisions = Division.objects.all()
    
    context = {
        'submissions': submissions,
        'divisions': divisions,
        'active_status_filter': status_filter,
        'active_division_filter': division_filter,
        'total_count': submissions.count(),
        'user_role': user.role.name if user.role else None,
    }
    
    return render(request, 'cbm/submissions.html', context)


@login_required
def cbm_create_call_view(request):
    """
    CBM procurement call creation view.
    Allows CBM users to create new procurement cycles with optional document upload.
    """
    user = request.user
    
    if user.role and user.role.name != 'CBM':
        return redirect('dashboard')
    
    from apps.procurement.models import ProcurementCall
    from datetime import datetime
    import uuid
    
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', '').strip()
        deadline_str = request.POST.get('deadline', '').strip()
        status = request.POST.get('status', 'Draft')
        call_document = request.FILES.get('call_document', None)
        
        # Validate required fields
        errors = []
        if not title:
            errors.append('Title is required.')
        if not description:
            errors.append('Description is required.')
        if not deadline_str:
            errors.append('Deadline is required.')
        
        if errors:
            context = {'error': ' '.join(errors)}
            return render(request, 'cbm/create_call.html', context)
        
        # Parse deadline from datetime-local format
        try:
            deadline = datetime.fromisoformat(deadline_str)
            deadline = timezone.make_aware(deadline)
        except (ValueError, TypeError) as e:
            context = {'error': f'Invalid deadline format: {str(e)}'}
            return render(request, 'cbm/create_call.html', context)
        
        # Validate document file size (max 10MB)
        if call_document and call_document.size > 10 * 1024 * 1024:
            context = {'error': 'Document file size must not exceed 10MB.'}
            return render(request, 'cbm/create_call.html', context)
        
        # Generate unique reference number
        reference_number = f"PC-{timezone.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8].upper()}"
        
        # Create the procurement call
        try:
            procurement_call = ProcurementCall.objects.create(
                title=title,
                description=description,
                reference_number=reference_number,
                start_date=timezone.now(),
                end_date=deadline,
                status=status,
                created_by=user,
                call_document=call_document,
            )
            # Redirect to calls list on success
            return redirect('cbm_calls')
        except Exception as e:
            import traceback
            error_msg = f'Error creating procurement call: {str(e)}'
            context = {'error': error_msg}
            return render(request, 'cbm/create_call.html', context)
    
    context = {}
    return render(request, 'cbm/create_call.html', context)


@login_required
def cbm_publish_call_view(request, call_id):
    """
    CBM publish procurement call view.
    Allows CBM users to publish draft procurement calls.
    """
    user = request.user
    
    if user.role and user.role.name != 'CBM':
        return redirect('dashboard')
    
    from apps.procurement.models import ProcurementCall
    
    if request.method == 'POST':
        try:
            call = ProcurementCall.objects.get(id=call_id, status='Draft')
            call.status = 'Active'
            call.save()
        except ProcurementCall.DoesNotExist:
            pass
    
    return redirect('cbm_calls')


@login_required
def cbm_call_detail_view(request, call_id):
    """
    CBM procurement call detail view.
    Allows CBM to view and edit procurement call details.
    """
    user = request.user
    
    if user.role and user.role.name != 'CBM':
        return redirect('dashboard')
    
    from apps.procurement.models import ProcurementCall
    from django.utils import timezone
    
    try:
        call = ProcurementCall.objects.get(id=call_id)
    except ProcurementCall.DoesNotExist:
        messages.error(request, 'Procurement call not found')
        return redirect('cbm_calls')
    
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', '').strip()
        deadline_str = request.POST.get('deadline', '').strip()
        call_document = request.FILES.get('call_document', None)
        
        # Validate required fields
        errors = []
        if not title:
            errors.append('Title is required.')
        if not description:
            errors.append('Description is required.')
        if not deadline_str:
            errors.append('Deadline is required.')
        
        if errors:
            context = {
                'call': call,
                'error': ' '.join(errors),
            }
            return render(request, 'cbm/call_detail.html', context)
        
        # Parse deadline
        try:
            deadline = timezone.datetime.fromisoformat(deadline_str)
            deadline = timezone.make_aware(deadline)
        except (ValueError, TypeError) as e:
            context = {
                'call': call,
                'error': f'Invalid deadline format',
            }
            return render(request, 'cbm/call_detail.html', context)
        
        # Validate document file size (max 10MB)
        if call_document and call_document.size > 10 * 1024 * 1024:
            context = {
                'call': call,
                'error': 'Document file size must not exceed 10MB.',
            }
            return render(request, 'cbm/call_detail.html', context)
        
        # Update call
        try:
            call.title = title
            call.description = description
            call.end_date = deadline
            if call_document:
                call.call_document = call_document
            call.save()
            messages.success(request, 'Procurement call updated successfully')
            return redirect('cbm_call_detail', call_id=call.id)
        except Exception as e:
            context = {
                'call': call,
                'error': f'Error updating call: {str(e)}',
            }
            return render(request, 'cbm/call_detail.html', context)
    
    # Get related submissions
    submissions = call.submissions.all().select_related('division')
    
    # Detect document file type for preview
    doc_file_type = None
    absolute_document_url = None
    doc_preview_html = None
    if call.call_document:
        filename = call.call_document.name.lower()
        if filename.endswith('.pdf'):
            doc_file_type = 'pdf'
        elif filename.endswith('.txt'):
            doc_file_type = 'txt'
        elif filename.endswith(('.doc', '.docx')):
            doc_file_type = 'word'
            # Convert Word to HTML
            try:
                doc_preview_html = convert_word_to_html(call.call_document.path)
            except Exception as e:
                doc_preview_html = f'<div class="alert alert-danger">Error loading Word document preview: {str(e)}</div>'
        elif filename.endswith(('.xls', '.xlsx')):
            doc_file_type = 'excel'
            # Convert Excel to HTML
            try:
                doc_preview_html = convert_excel_to_html(call.call_document.path)
            except Exception as e:
                doc_preview_html = f'<div class="alert alert-danger">Error loading Excel file preview: {str(e)}</div>'
        elif filename.endswith('.zip'):
            doc_file_type = 'zip'
        # Convert relative URL to absolute URL for Google Docs Viewer
        absolute_document_url = request.build_absolute_uri(call.call_document.url)
    
    context = {
        'call': call,
        'submissions': submissions,
        'total_submissions': call.submissions.count(),
        'doc_file_type': doc_file_type,
        'absolute_document_url': absolute_document_url,
        'doc_preview_html': doc_preview_html,
    }
    return render(request, 'cbm/call_detail.html', context)


@login_required
def cbm_submission_detail_view(request, submission_id):
    """
    CBM submission detail view with 9-stage workflow tracker.
    Shows workflow history, current stage, and approval actions.
    """
    user = request.user
    
    if user.role and user.role.name != 'CBM':
        return redirect('dashboard')
    
    from apps.procurement.models import Submission
    from apps.workflows.models import WorkflowHistory
    from datetime import timedelta
    
    try:
        submission = Submission.objects.select_related(
            'call', 'division', 'current_stage', 'submitted_by'
        ).get(id=submission_id)
    except Submission.DoesNotExist:
        return redirect('cbm_submissions')
    
    # Get workflow history for this submission
    workflow_history = WorkflowHistory.objects.filter(
        submission=submission
    ).select_related('from_stage', 'to_stage').order_by('created_at')
    
    # Calculate progress metrics
    start_date = submission.created_at
    end_date = submission.call.end_date if submission.call else timezone.now()
    days_elapsed = (timezone.now() - start_date).days
    days_remaining = max(0, (end_date - timezone.now()).days)
    total_days = (end_date - start_date).days
    
    # Define workflow stages in order with role assignment (20 stages total)
    # These are the 20 procurement stages CBM oversees/monitors
    workflow_stages = [
        {'index': 1, 'name': 'Call Issued', 'role': 'CBM', 'responsible': 'Chief Budget Manager'},
        {'index': 2, 'name': 'HOD/DM Submit', 'role': 'HOD/DM', 'responsible': 'Head of Department'},
        {'index': 3, 'name': 'Review of Procurement Draft', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 4, 'name': 'Submit Compiled Document', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 5, 'name': 'CBM Review', 'role': 'CBM', 'responsible': 'Chief Budget Manager'},
        {'index': 6, 'name': 'Publish Plan', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 7, 'name': 'Prepare Tender Document', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 8, 'name': 'CBM Review TD', 'role': 'CBM', 'responsible': 'Chief Budget Manager'},
        {'index': 9, 'name': 'Publication of TD', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 10, 'name': 'Opening', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 11, 'name': 'Evaluation', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 12, 'name': 'CBM Approval', 'role': 'CBM', 'responsible': 'Chief Budget Manager'},
        {'index': 13, 'name': 'Notify Bidders', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 14, 'name': 'Contract Negotiation', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 15, 'name': 'Contract Drafting', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 16, 'name': 'Legal Review', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 17, 'name': 'Supplier Approval', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 18, 'name': 'MINIJUST Legal Review', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 19, 'name': 'Awarded', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
        {'index': 20, 'name': 'Completed', 'role': 'Procurement Team', 'responsible': 'Procurement Officer'},
    ]

    # Map submission status to workflow stage index
    status_to_stage_index = {
        'Draft': 0,  # Pre-workflow (division user creating request)
        'Call Issued': 1,
        'HOD/DM Submit': 2,
        'Review of Procurement Draft': 3,
        'Returned': -1,  # Special handling for returned status
        'Submit Compiled Document': 4,
        'CBM Review': 5,
        'Publish Plan': 6,
        'Prepare Tender Document': 7,
        'CBM Review TD': 8,
        'Publication of TD': 9,
        'Opening': 10,
        'Evaluation': 11,
        'CBM Approval': 12,
        'Notify Bidders': 13,
        'Contract Negotiation': 14,
        'Contract Drafting': 15,
        'Legal Review': 16,
        'Supplier Approval': 17,
        'MINIJUST Legal Review': 18,
        'Awarded': 19,
        'Completed': 20,
        'Rejected': 5,  # Rejected at CBM review level
        'Cancelled': 5,  # Cancelled at CBM review level
    }
    
    # Get the actual stage index from submission's current_stage (the single source of truth)
    # This ensures that when clarification updates current_stage, the view reflects the new stage
    if submission.current_stage:
        current_stage_index = submission.current_stage.order
        is_returned = submission.status == 'Returned'
    else:
        current_stage_index = status_to_stage_index.get(submission.status, 0)
        is_returned = submission.status == 'Returned'
    
    # Calculate progress percentage based on stage progress (not time)
    progress_percentage = int((current_stage_index / 20) * 100) if current_stage_index > 0 else 0
    
    # Build stage progression data with role information
    stages_progress = []
    for stage_info in workflow_stages:
        idx = stage_info['index']
        
        # Get approval date from workflow history for this stage
        approval_date = None
        if idx >= 6:  # Only from Publish Plan (stage 6) onwards
            # Look for history where FROM_STAGE is this stage (the stage they were at when they approved)
            history_record = workflow_history.filter(from_stage__order=idx).first()
            if history_record and history_record.approval_date:
                approval_date = history_record.approval_date
        
        stages_progress.append({
            'index': idx,
            'order': idx,
            'name': stage_info['name'],
            'role': stage_info['role'],
            'responsible': stage_info['responsible'],
            'completed': idx < current_stage_index,
            'current': idx == current_stage_index,
            'pending': idx > current_stage_index,
            'approval_date': approval_date,
        })
    
    # Split stages_progress into two separate tracker groups
    submission_workflow_viz_cbm = [s for s in stages_progress if s['order'] <= 3]
    compiled_workflow_viz_cbm = [s for s in stages_progress if 4 <= s['order'] <= 6]

    # Get supporting documents (attachments from submission)
    supporting_documents = []
    if submission.attachments:
        supporting_documents = submission.attachments if isinstance(submission.attachments, list) else []
    
    # Get tenders if submission is Plan Published
    from apps.procurement.models import Tender
    tenders = Tender.objects.filter(submission=submission).order_by('-created_at') if submission.status == 'Plan Published' else []

    context = {
        'submission': submission,
        'workflow_history': workflow_history,
        'workflow_stages': stages_progress,
        'submission_workflow_viz': submission_workflow_viz_cbm,
        'compiled_workflow_viz': compiled_workflow_viz_cbm,
        'days_elapsed': days_elapsed,
        'days_remaining': days_remaining,
        'progress_percentage': progress_percentage,
        'is_returned': is_returned,
        'returned_status': 'Returned for Clarification' if is_returned else None,
        'supporting_documents': supporting_documents,
        'tenders': tenders,
    }
    
    return render(request, 'cbm/submission_detail.html', context)


@login_required
def cbm_approve_submission_view(request, submission_id):
    """
    Approve a submission based on current stage and move to next stage.
    - If at CBM Review (stage 5): Move to Publish Plan (stage 6)
    - If at CBM Review TD (stage 8): Move to Publication of TD (stage 9)
    - If at Legal Review (stage 16): Move to MINIJUST Legal Review (stage 18) or Awarded (stage 19)

    For stages requiring procurement method/tender type, prompts for selection first.
    """
    user = request.user

    # Allow CBM and Procurement Team
    if user.role and user.role.name not in ['CBM', 'Procurement Team']:
        return redirect('dashboard')

    # Procurement Team can only access at Step 7 (Prepare Tender Document)
    from apps.procurement.models import Submission
    try:
        submission = Submission.objects.get(id=submission_id)
        current_stage = submission.current_stage
        if user.role and user.role.name == 'Procurement Team':
            if not current_stage or current_stage.order != 7:
                return redirect('procurement_submissions')
    except Submission.DoesNotExist:
        pass
    
    from apps.workflows.models import WorkflowHistory, WorkflowStage
    from apps.notifications.models import Notification
    from apps.procurement.forms import ProcurementMethodForm, TenderTypeForm
    from django.contrib import messages
    
    try:
        submission = Submission.objects.get(id=submission_id)
    except Submission.DoesNotExist:
        if user.role and user.role.name == 'Procurement Team':
            messages.error(request, 'Submission not found')
            return redirect('procurement_submissions')
        else:
            messages.error(request, 'Submission not found')
            return redirect('cbm_submissions')
    
    current_stage = submission.current_stage
    
    # Determine next stage based on current stage
    if not current_stage:
        messages.error(request, 'Current stage not set')
        return redirect('cbm_submission_detail', submission_id=submission_id)
    
    # CBM handles stages 5, 8, 12, 13, 16, 17, 18, 19
    next_stage = None
    next_status = None
    requires_method_selection = False
    requires_tender_type = False

    if current_stage.order == 5:  # CBM Review → Publish Plan
        next_stage = WorkflowStage.objects.filter(order=6).first()
        next_status = 'Publish Plan'
    elif current_stage.order == 6:  # Publish Plan → Prepare Tender Document
        next_stage = WorkflowStage.objects.filter(order=7).first()
        next_status = 'Prepare Tender Document'
    elif current_stage.order == 7:  # Prepare Tender Document - Show procurement method form AFTER approval
        next_stage = WorkflowStage.objects.filter(order=8).first()
        next_status = 'CBM Review TD'
        # Set flag to show method form instead of advancing immediately
        requires_method_selection = True
    elif current_stage.order == 8:  # CBM Review TD → Publication of TD
        next_stage = WorkflowStage.objects.filter(order=9).first()
        next_status = 'Publication of TD'
    elif current_stage.order == 9:  # Publication of TD → Opening
        next_stage = WorkflowStage.objects.filter(order=10).first()
        next_status = 'Opening'
    elif current_stage.order == 10:  # Opening → Evaluation
        next_stage = WorkflowStage.objects.filter(order=11).first()
        next_status = 'Evaluation'
    elif current_stage.order == 11:  # Evaluation → CBM Approval
        next_stage = WorkflowStage.objects.filter(order=12).first()
        next_status = 'CBM Approval'
    elif current_stage.order == 12:  # CBM Approval → Notify Bidders
        next_stage = WorkflowStage.objects.filter(order=13).first()
        next_status = 'Notify Bidders'
    elif current_stage.order == 13:  # Notify Bidders → Contract Negotiation
        next_stage = WorkflowStage.objects.filter(order=14).first()
        next_status = 'Contract Negotiation'
    elif current_stage.order == 14:  # Contract Negotiation → Contract Drafting
        next_stage = WorkflowStage.objects.filter(order=15).first()
        next_status = 'Contract Drafting'
    elif current_stage.order == 15:  # Contract Drafting → Legal Review
        next_stage = WorkflowStage.objects.filter(order=16).first()
        next_status = 'Legal Review'
    elif current_stage.order == 16:  # Legal Review → Supplier Approval
        next_stage = WorkflowStage.objects.filter(order=17).first()
        next_status = 'Supplier Approval'
    elif current_stage.order == 17:  # Supplier Approval
        # Check if amount requires MINIJUST review (> 500M RWF)
        if submission.total_budget and submission.total_budget > 500000000:
            next_stage = WorkflowStage.objects.filter(order=18).first()
            next_status = 'MINIJUST Legal Review'
        else:
            next_stage = WorkflowStage.objects.filter(order=19).first()
            next_status = 'Awarded'
    elif current_stage.order == 18:  # MINIJUST Legal Review → Awarded
        next_stage = WorkflowStage.objects.filter(order=19).first()
        next_status = 'Awarded'
    elif current_stage.order == 19:  # Awarded → Completed
        next_stage = WorkflowStage.objects.filter(order=20).first()
        next_status = 'Completed'
    else:
        messages.error(request, f'CBM cannot approve at stage {current_stage.name}')
        return redirect('cbm_submission_detail', submission_id=submission_id)
    
    # Handle GET request - show form if method selection needed
    if request.method == 'GET' and requires_method_selection:
        form = ProcurementMethodForm()
        context = {
            'submission': submission,
            'form': form,
            'next_status': next_status,
            'requires_method_selection': True,
        }
        return render(request, 'dashboard/cbm_approve_with_method.html', context)
    
    # Handle POST request
    if request.method == 'POST':
        notes = request.POST.get('notes', '')
        approval_date_str = request.POST.get('approval_date', '')
        approval_date = None
        
        # Parse approval_date if provided
        if approval_date_str:
            try:
                from datetime import datetime
                approval_date = datetime.strptime(approval_date_str, '%Y-%m-%d').date()
            except ValueError:
                messages.warning(request, 'Invalid approval date format')
        
        # Get method/tender selection if needed
        procurement_method = None
        tender_type = None
        
        # For Step 7: First POST from modal shows form, second POST from form processes method
        if current_stage.order == 7 and requires_method_selection:
            # Check if this is from method form or from modal
            if 'procurement_method' in request.POST:
                # This is the method form submission - process it
                form = ProcurementMethodForm(request.POST)
                if not form.is_valid():
                    context = {
                        'submission': submission,
                        'form': form,
                        'next_status': next_status,
                        'requires_method_selection': True,
                    }
                    return render(request, 'dashboard/cbm_approve_with_method.html', context)
                procurement_method = form.cleaned_data['procurement_method']
            else:
                # This is the modal submission - show the method form
                form = ProcurementMethodForm()
                context = {
                    'submission': submission,
                    'form': form,
                    'next_status': next_status,
                    'requires_method_selection': True,
                }
                return render(request, 'dashboard/cbm_approve_with_method.html', context)
        elif requires_method_selection:
            form = ProcurementMethodForm(request.POST)
            if not form.is_valid():
                context = {
                    'submission': submission,
                    'form': form,
                    'next_status': next_status,
                    'requires_method_selection': True,
                }
                return render(request, 'dashboard/cbm_approve_with_method.html', context)
            procurement_method = form.cleaned_data['procurement_method']
        
        # Update submission
        submission.status = next_status
        if next_stage:
            submission.current_stage = next_stage
        
        # Save procurement method if selected at Publish Plan stage
        if requires_method_selection and procurement_method:
            submission.procurement_method = procurement_method
        
        submission.save()
        
        # Set timeline if next stage requires it
        if next_status == 'Prepare Tender Document':
            # Timeline info is added at Step 6, using the procurement method selected at Step 5
            submission.set_timeline_deadline(
                stage_name='Publication of TD',
                procurement_method=submission.procurement_method
            )
        elif next_status in ['Publication of TD']:
            submission.set_timeline_deadline(
                stage_name=next_status,
                procurement_method=submission.procurement_method or procurement_method
            )
        elif next_status in ['Evaluation']:
            submission.set_timeline_deadline(
                stage_name=next_status,
                procurement_method=None
            )
        elif next_status in ['Notification']:
            submission.set_timeline_deadline(
                stage_name=next_status,
                procurement_method=None
            )
        # Note: Bid Validity and Contract Signature handled by Procurement Team
        
        # Record in workflow history with approval_date if next stage is Publish Plan or later
        history_data = {
            'submission': submission,
            'from_stage': current_stage,
            'to_stage': next_stage,
            'action': 'approve',
            'comments': notes or f'Approved by CBM - Moving to {next_status}',
            'action_by': user,
        }
        
        # Only set approval_date for stages from Publish Plan (order 6) onwards
        if next_stage and next_stage.order >= 6 and approval_date:
            history_data['approval_date'] = approval_date
        
        WorkflowHistory.objects.create(**history_data)
        
        # Notify the division
        if submission.submitted_by:
            Notification.objects.create(
                user=submission.submitted_by,
                title=f'Submission Approved: {submission.tracking_reference}',
                message=f'Your procurement request for {submission.item_name} has been approved by CBM and moved to {next_status}.',
                notification_type='submission_status',
                priority='high',
                related_object_type='Submission',
                related_object_id=str(submission.id),
                action_url=f'/dashboard/hod/submissions/{submission.id}/',
            )
        
        # Notify Procurement Team if moving to Publish Plan or Publication of TD
        if next_status in ['Publish Plan', 'Publication of TD']:
            from apps.accounts.models import User
            procurement_users = User.objects.filter(role__name='Procurement Team', is_active=True)
            for proc_user in procurement_users:
                Notification.objects.create(
                    user=proc_user,
                    title=f'Submission Ready: {submission.tracking_reference}',
                    message=f'Submission {submission.tracking_reference} is ready for {next_status}. Please proceed with external activities.',
                    notification_type='submission_status',
                    priority='high',
                    related_object_type='Submission',
                    related_object_id=str(submission.id),
                    action_url=f'/dashboard/procurement/submissions/{submission.id}/',
                )
        
        messages.success(request, f'Submission approved and moved to {next_status}!')
        
        # Redirect to appropriate view based on user role
        if user.role and user.role.name == 'Procurement Team':
            return redirect('procurement_submission_detail', submission_id=submission_id)
        else:
            return redirect('cbm_submission_detail', submission_id=submission_id)
    
    # Default GET (without method selection needed)
    # Redirect to POST to proceed with approval
    return redirect('cbm_submission_detail', submission_id=submission_id)


@login_required
def cbm_request_clarification_view(request, submission_id):
    """
    Request clarification on a submission.
    Moves submission back one stage and sets status to Returned.
    """
    user = request.user
    
    if user.role and user.role.name != 'CBM':
        return redirect('dashboard')
    
    from apps.procurement.models import Submission
    from apps.workflows.models import WorkflowHistory, WorkflowStage
    from apps.notifications.models import Notification
    from django.contrib import messages
    
    if request.method == 'POST':
        try:
            submission = Submission.objects.get(id=submission_id)
            clarification_request = request.POST.get('clarification_request', '')
            
            current_stage = submission.current_stage
            
            if not current_stage:
                messages.error(request, 'Current stage not set')
                return redirect('cbm_submission_detail', submission_id=submission_id)
            
            # Get the previous stage (order - 1)
            previous_stage = WorkflowStage.objects.filter(order=current_stage.order - 1).first()
            
            if not previous_stage:
                messages.error(request, f'Cannot return from the first stage')
                return redirect('cbm_submission_detail', submission_id=submission_id)
            
            # Update submission to previous stage with Returned status
            submission.current_stage = previous_stage
            submission.status = 'Returned'
            submission.save()
            
            # Record in workflow history
            WorkflowHistory.objects.create(
                submission=submission,
                from_stage=current_stage,
                to_stage=previous_stage,
                action='return',
                comments=clarification_request or f'Returned for clarification by CBM',
                action_by=user,
            )
            
            # Notify the division that clarification is requested
            if submission.submitted_by:
                Notification.objects.create(
                    user=submission.submitted_by,
                    title=f'Clarification Requested: {submission.tracking_reference}',
                    message=f'CBM has requested clarifications for {submission.item_name}. Please review and resubmit.',
                    notification_type='approval_required',
                    priority='high',
                    related_object_type='Submission',
                    related_object_id=str(submission.id),
                    action_url=f'/dashboard/hod/submissions/{submission.id}/',
                )
            
            messages.success(request, 'Submission returned for clarification!')
            
        except Submission.DoesNotExist:
            messages.error(request, 'Submission not found')
        except Exception as e:
            messages.error(request, f'Error: {str(e)}')
    
    return redirect('cbm_submission_detail', submission_id=submission_id)


@login_required
def cbm_reject_submission_view(request, submission_id):
    """Reject a submission at CBM stage."""
    user = request.user
    
    if user.role and user.role.name != 'CBM':
        return redirect('dashboard')
    
    from apps.procurement.models import Submission
    from apps.workflows.models import WorkflowHistory
    from apps.notifications.models import Notification
    
    if request.method == 'POST':
        try:
            submission = Submission.objects.get(id=submission_id)
            rejection_reason = request.POST.get('rejection_reason', '')
            
            # Mark as Rejected
            submission.status = 'Rejected'
            submission.save()
            
            # Record in workflow history
            WorkflowHistory.objects.create(
                submission=submission,
                from_stage=submission.current_stage,
                to_stage=submission.current_stage,
                action='reject',
                comments=rejection_reason or 'Rejected by CBM',
                action_by=user,
            )
            
            # Notify the division that submission was rejected
            if submission.submitted_by:
                Notification.objects.create(
                    user=submission.submitted_by,
                    title=f'Submission Rejected: {submission.tracking_reference}',
                    message=f'Your procurement request for {submission.item_name} has been rejected by CBM.',
                    notification_type='submission_status',
                    priority='high',
                    related_object_type='Submission',
                    related_object_id=str(submission.id),
                    action_url=f'/dashboard/hod/submissions/{submission.id}/',
                )
            
        except Submission.DoesNotExist:
            pass
    
    return redirect('cbm_submission_detail', submission_id=submission_id)


@login_required
def cbm_publish_submission_view(request, submission_id):
    """Publish a submission to Umucyo."""
    user = request.user
    
    if user.role and user.role.name != 'CBM':
        return redirect('dashboard')
    
    from apps.procurement.models import Submission
    from apps.workflows.models import WorkflowHistory
    
    if request.method == 'POST':
        try:
            submission = Submission.objects.get(id=submission_id)
            
            # Update submission status
            submission.status = 'Published to Umucyo'
            submission.save()
            
            # Record in workflow history
            WorkflowHistory.objects.create(
                submission=submission,
                action='publish',
                comments='Published to Umucyo platform',
                action_by=user,
            )
            
        except Submission.DoesNotExist:
            pass
    
    return redirect('cbm_submission_detail', submission_id=submission_id)


# ============================================================================
# HOD / DIVISION MANAGER VIEWS
# ============================================================================

@login_required
def hod_dashboard_view(request):
    """
    HOD/Division Manager dashboard.
    Shows division-specific KPIs, active calls, submissions, and tracking.
    """
    user = request.user

    # Only HOD/DM and Admin can access
    if user.role and user.role.name not in ['HOD/DM', 'Admin']:
        return redirect('dashboard')

    from apps.procurement.models import ProcurementCall, Submission
    from apps.workflows.models import WorkflowStage, WorkflowHistory
    from apps.notifications.models import Notification
    from django.db.models import Sum

    # Get division
    division = user.division
    if not division:
        context = {'error': 'No division assigned to your account'}
        return render(request, 'dashboard/hod_dashboard.html', context)

    # Get active calls
    active_calls = ProcurementCall.objects.filter(status='Active')

    # Get division's submissions
    submissions = Submission.objects.filter(
        division=division,
        is_deleted=False
    ).select_related('call', 'current_stage')

    # ── Pipeline stage buckets ────────────────────────────────────
    draft_statuses      = ['Draft', 'Call Issued']
    submitted_statuses  = ['HOD/DM Submit']
    review_statuses     = ['Review of Procurement Draft', 'Submit Compiled Document',
                           'CBM Review', 'Publish Plan']
    tender_statuses     = ['Prepare Tender Document', 'CBM Review TD', 'Publication of TD',
                           'Opening', 'Evaluation', 'CBM Approval', 'Notify Bidders',
                           'Contract Negotiation', 'Contract Drafting', 'Legal Review',
                           'Supplier Approval', 'MINIJUST Legal Review']
    final_statuses      = ['Awarded', 'Completed']

    pipeline = {
        'draft':      submissions.filter(status__in=draft_statuses).count(),
        'submitted':  submissions.filter(status__in=submitted_statuses).count(),
        'review':     submissions.filter(status__in=review_statuses).count(),
        'tender':     submissions.filter(status__in=tender_statuses).count(),
        'awarded':    submissions.filter(status__in=final_statuses).count(),
        'returned':   submissions.filter(status='Returned').count(),
    }

    # ── Budget breakdown ──────────────────────────────────────────
    budget_total     = submissions.aggregate(t=Sum('total_budget'))['t'] or 0
    budget_submitted = submissions.filter(
        status__in=submitted_statuses + review_statuses + tender_statuses
    ).aggregate(t=Sum('total_budget'))['t'] or 0
    budget_awarded   = submissions.filter(
        status__in=final_statuses
    ).aggregate(t=Sum('total_budget'))['t'] or 0

    # ── Urgent: returned submissions ──────────────────────────────
    urgent_returned = []
    for sub in submissions.filter(status='Returned').select_related('call').order_by('-updated_at'):
        urgent_returned.append({
            'id': sub.id,
            'tracking_reference': sub.tracking_reference,
            'item_name': sub.item_name,
            'call_ref': sub.call.reference_number if sub.call else '—',
            'days_since_returned': (timezone.now() - sub.updated_at).days,
        })

    # ── Calculate KPI stats ───────────────────────────────────────
    stats = {
        'division_name': division.name,
        'active_calls': active_calls.count(),
        'total_submissions': submissions.count(),
        'returned': pipeline['returned'],
        'in_progress': pipeline['submitted'] + pipeline['review'] + pipeline['tender'],
        'awarded_completed': pipeline['awarded'],
        'total_budget': budget_total,
        # kept for backward compat
        'pending_approvals': submissions.filter(
            status__in=['Submitted', 'Under Review']
        ).count(),
        'approved_count': submissions.filter(status='Approved').count(),
        'clarification_needed': pipeline['returned'],
        'completed': submissions.filter(status='Completed').count(),
    }

    # ── Get submission breakdown by status ────────────────────────
    submission_breakdown = []
    for status, _ in Submission.STATUS_CHOICES:
        count = submissions.filter(status=status).count()
        if count > 0:
            status_class = get_status_class(status)
            submission_breakdown.append({
                'status': status,
                'count': count,
                'class': status_class,
            })

    # ── Active calls with submission deadline info ──────────────
    active_calls_list = []
    for call in active_calls:
        division_submission = submissions.filter(call=call).first()
        days_remaining = call.days_remaining
        deadline_class = 'danger' if days_remaining <= 3 else ('warning' if days_remaining <= 7 else 'info')
        active_calls_list.append({
            'id': call.id,
            'reference_number': call.reference_number,
            'title': call.title,
            'start_date': call.start_date,
            'end_date': call.end_date,
            'days_remaining': days_remaining,
            'deadline_class': deadline_class,
            'has_submission': division_submission is not None,
            'submission_id': division_submission.id if division_submission else None,
            'submission_status': division_submission.status if division_submission else None,
        })

    # ── Recent submissions with workflow progress ─────────────────
    recent_submissions = []
    recent_subs = submissions.order_by('-created_at')[:8]
    for sub in recent_subs:
        workflow_history = WorkflowHistory.objects.filter(submission=sub).order_by('created_at')
        stages_completed = workflow_history.values('from_stage').distinct().count()
        total_stages = WorkflowStage.objects.count()
        progress_percent = (stages_completed / total_stages * 100) if total_stages > 0 else 0
        recent_submissions.append({
            'id': sub.id,
            'tracking_reference': sub.tracking_reference,
            'item_name': sub.item_name,
            'call_ref': sub.call.reference_number if sub.call else '—',
            'quantity': sub.quantity,
            'total_budget': sub.total_budget,
            'status': sub.status,
            'status_class': get_status_class(sub.status),
            'current_stage': sub.current_stage.name if sub.current_stage else 'Not Started',
            'progress_percent': int(progress_percent),
            'created_at': sub.created_at,
            'days_at_stage': sub.days_at_current_stage,
        })

    # ── Unread notifications ──────────────────────────────────────
    unread_notifications = Notification.objects.filter(
        user=user, is_read=False
    ).select_related().order_by('-created_at')[:5]
    unread_count = Notification.objects.filter(user=user, is_read=False).count()

    context = {
        'stats': stats,
        'pipeline': pipeline,
        'budget_total': budget_total,
        'budget_submitted': budget_submitted,
        'budget_awarded': budget_awarded,
        'urgent_returned': urgent_returned,
        'submission_breakdown': submission_breakdown,
        'recent_submissions': recent_submissions,
        'active_calls_list': active_calls_list,
        'unread_notifications': unread_notifications,
        'unread_count': unread_count,
        'division': division,
    }

    return render(request, 'dashboard/hod_dashboard.html', context)


@login_required
def hod_submission_list_view(request):
    """
    View list of all division submissions with filtering and sorting.
    """
    user = request.user
    
    if user.role and user.role.name not in ['HOD/DM', 'Admin']:
        return redirect('dashboard')
    
    from apps.procurement.models import Submission, ProcurementCall
    
    division = user.division
    if not division:
        context = {'error': 'No division assigned'}
        return render(request, 'dashboard/hod_submissions.html', context)
    
    submissions = Submission.objects.filter(
        division=division,
        is_deleted=False
    ).select_related('call', 'current_stage').order_by('-created_at')
    
    # Calculate KPI statistics
    stats = {
        'total_submissions': submissions.count(),
        'draft': submissions.filter(status='Draft').count(),
        'pending_review': submissions.filter(status__in=['HOD/DM Submit', 'Review of Procurement Draft']).count(),
        'with_procurement': submissions.filter(status__in=['Submit Compiled Document', 'CBM Review', 'Publish Plan']).count(),
        'under_cbm_review': submissions.filter(status__in=['CBM Review', 'CBM Review TD']).count(),
        'approved': submissions.filter(status__in=['Publish Plan', 'Prepare Tender Document', 'Publication of TD']).count(),
        'returned': submissions.filter(status='Returned').count(),
        'awarded': submissions.filter(status='Awarded').count(),
    }
    
    # Filter by status if provided
    status_filter = request.GET.get('status')
    if status_filter:
        submissions = submissions.filter(status=status_filter)
    
    # Filter by call if provided
    call_filter = request.GET.get('call')
    if call_filter:
        submissions = submissions.filter(call__id=call_filter)
    
    # Get all available calls for filter dropdown
    calls = ProcurementCall.objects.filter(status='Active')
    
    # Get status choices for filter
    status_choices = [status for status, _ in Submission.STATUS_CHOICES]
    
    # Check if call_id is provided (from notification click)
    preselect_call_id = request.GET.get('call_id')
    
    context = {
        'submissions': submissions,
        'calls': calls,
        'status_choices': status_choices,
        'division': division,
        'selected_status': status_filter,
        'selected_call': call_filter,
        'preselect_call_id': preselect_call_id,
        'stats': stats,
        'total_count': Submission.objects.filter(division=division, is_deleted=False).count(),
    }
    
    return render(request, 'dashboard/hod_submissions.html', context)


@login_required
def hod_submission_detail_view(request, submission_id):
    """
    View detailed submission with full workflow tracking and comments.
    """
    user = request.user
    
    if user.role and user.role.name not in ['HOD/DM', 'Admin']:
        return redirect('dashboard')
    
    from apps.procurement.models import Submission, Comment
    from apps.workflows.models import WorkflowHistory, WorkflowStage
    
    try:
        submission = Submission.objects.select_related(
            'division', 'call', 'current_stage', 'submitted_by'
        ).get(id=submission_id)
    except Submission.DoesNotExist:
        return render(request, 'dashboard/error.html', {'error': 'Submission not found'})
    
    # Check permission - user can only view their division's submissions
    if user.role and user.role.name == 'HOD/DM' and user.division != submission.division:
        return render(request, 'dashboard/error.html', {'error': 'Access denied'})
    
    # Get workflow history
    workflow_history = WorkflowHistory.objects.filter(
        submission=submission
    ).select_related('action_by', 'from_stage', 'to_stage').order_by('created_at')
    
    # Build workflow visualization using the same logic as CBM
    # HOD/DM only sees internal stages (1-3); beyond that shows as 'With Procurement'
    # HOD/DM only sees submission stages (1-3); Compiled Document tracker is for Procurement Team/CBM
    all_stages = WorkflowStage.objects.filter(order__lte=3).order_by('order')
    
    # Map submission status to workflow stage order
    status_to_stage_order = {
        'Draft': 0,  # Pre-workflow
        'Call Issued': 1,
        'HOD/DM Submit': 2,
        'Review of Procurement Draft': 3,  # Next stage after HOD/DM Submit
        'Returned': -1,  # Special handling
        'Submit Compiled Document': 4,
        'CBM Review': 5,
        'Publish Plan': 6,
        'Prepare Tender Document': 7,
        'CBM Review TD': 8,
        'Publication of TD': 9,
        'Opening': 10,
        'Evaluation': 11,
        'CBM Approval': 12,
        'Notify Bidders': 13,
        'Contract Negotiation': 14,
        'Contract Drafting': 15,
        'Legal Review': 16,
        'Supplier Approval': 17,
        'MINIJUST Legal Review': 18,
        'Awarded': 19,
        'Completed': 20,
        'Rejected': 5,
        'Cancelled': 5,
    }
    
    # Get the current stage order from submission's current_stage (the single source of truth)
    # This ensures that when clarification updates current_stage, the view reflects the new stage
    if submission.current_stage:
        current_stage_order = submission.current_stage.order
        is_returned = submission.status == 'Returned'
    else:
        # Fallback to status mapping if current_stage is not set
        current_stage_order = status_to_stage_order.get(submission.status, 0)
        is_returned = submission.status == 'Returned'
    
    # Build workflow visualization with correct stage states
    workflow_visualization = []
    completed_count = 0
    
    for stage in all_stages:
        # A stage is:
        # - Completed: if its order is less than current_stage_order
        # - Current: if its order equals current_stage_order
        # - Pending: if its order is greater than current_stage_order
        is_completed = stage.order < current_stage_order
        is_current = stage.order == current_stage_order
        
        # Get history entry for this stage (for display purposes)
        history_entry = workflow_history.filter(to_stage=stage).first()
        
        if is_completed:
            completed_count += 1
        
        # Get approval date from history entry if available
        # Look for records where FROM_STAGE is this stage (when this stage was approved and moved forward)
        approval_date = None
        if stage.order >= 5:  # Only from CBM Review onwards (order 5)
            history_record = workflow_history.filter(from_stage=stage).first()
            if history_record and history_record.approval_date:
                approval_date = history_record.approval_date
        
        workflow_visualization.append({
            'id': stage.id,
            'order': stage.order,
            'name': stage.name,
            'description': stage.description,
            'color': stage.color,
            'icon': stage.icon,
            'completed': is_completed,
            'is_current': is_current,
            'completed_at': history_entry.created_at if history_entry else None,
            'completed_by': history_entry.action_by.full_name if history_entry and history_entry.action_by else None,
            'action_description': history_entry.comments if history_entry else None,
            'approval_date': approval_date,
        })
    
    # Calculate progress
    total_stages = all_stages.count()
    progress_percent = int((completed_count / total_stages * 100)) if total_stages > 0 else 0
    
    # Calculate SVG line x2 position (50 to 950, with 900 units of width)
    progress_x2 = 50 + (completed_count / total_stages * 900) if total_stages > 0 else 50
    
    # Calculate days elapsed and remaining
    from datetime import datetime
    now = timezone.now()
    days_elapsed = (now - submission.created_at).days
    days_remaining = 0
    if submission.expected_delivery_date:
        # expected_delivery_date is already a date object
        days_remaining = max(0, (submission.expected_delivery_date - now.date()).days)
    
    # Get comments/clarifications
    comments = Comment.objects.filter(
        submission=submission
    ).select_related('author').order_by('-created_at')
    
    # Get supporting documents grouped by type
    from apps.procurement.models import SubmissionDocument
    supporting_documents = submission.supporting_documents.all().order_by('document_type', '-uploaded_at')
    
    # Group documents by type and prepare preview data
    procurement_plan_docs = []
    technical_spec_docs = []
    market_survey_docs = []
    
    doc_type_groups = {
        'procurement_plan': procurement_plan_docs,
        'technical_specification': technical_spec_docs,
        'market_survey': market_survey_docs,
    }
    
    for doc in supporting_documents:
        # Detect document file type and prepare preview
        filename = doc.original_filename.lower()
        doc_file_type = None
        doc_preview_html = None
        absolute_document_url = request.build_absolute_uri(doc.file.url)
        
        if filename.endswith('.pdf'):
            doc_file_type = 'pdf'
        elif filename.endswith('.txt'):
            doc_file_type = 'txt'
        elif filename.endswith(('.doc', '.docx')):
            doc_file_type = 'word'
            # Convert Word to HTML
            try:
                doc_preview_html = convert_word_to_html(doc.file.path)
            except Exception as e:
                doc_preview_html = f'<div class="alert alert-danger">Error loading Word document preview: {str(e)}</div>'
        elif filename.endswith(('.xls', '.xlsx')):
            doc_file_type = 'excel'
            # Convert Excel to HTML
            try:
                doc_preview_html = convert_excel_to_html(doc.file.path)
            except Exception as e:
                doc_preview_html = f'<div class="alert alert-danger">Error loading Excel file preview: {str(e)}</div>'
        elif filename.endswith('.zip'):
            doc_file_type = 'zip'
        
        # Add document with preview data to the appropriate category
        doc_data = {
            'id': str(doc.id),
            'original_filename': doc.original_filename,
            'file_size': doc.file_size,
            'file_url': doc.file.url,
            'absolute_url': absolute_document_url,
            'doc_file_type': doc_file_type,
            'doc_preview_html': doc_preview_html,
            'uploaded_at': doc.uploaded_at,
            'uploaded_by': doc.uploaded_by.full_name if doc.uploaded_by else 'System',
        }
        
        if doc.document_type in doc_type_groups:
            doc_type_groups[doc.document_type].append(doc_data)
    
    # Split workflow visualization into two separate tracker groups
    submission_workflow_viz = [s for s in workflow_visualization if s['order'] <= 3]
    compiled_workflow_viz = [s for s in workflow_visualization if 4 <= s['order'] <= 6]

    # Determine if submission is editable
    is_editable = submission.status in ['Draft', 'Returned']
    
    # Get previous stage
    previous_stage = submission.current_stage.previous_stage if submission.current_stage else None
    
    context = {
        'submission': submission,
        'workflow_visualization': workflow_visualization,
        'submission_workflow_viz': submission_workflow_viz,
        'compiled_workflow_viz': compiled_workflow_viz,
        'workflow_timeline': workflow_history,
        'workflow_history': workflow_history,
        'comments': comments,
        'supporting_documents': supporting_documents,
        'procurement_plan_docs': procurement_plan_docs,
        'technical_spec_docs': technical_spec_docs,
        'market_survey_docs': market_survey_docs,
        'is_editable': is_editable,
        'previous_stage': previous_stage,
        'division': user.division,
        'current_stage_order': current_stage_order,
        'completed_count': completed_count,
        'total_stages': total_stages,
        'progress_percent': progress_percent,
        'progress_x2': progress_x2,
        'elapsed_days': days_elapsed,
        'remaining_days': days_remaining,
        'is_returned': is_returned,
    }
    
    return render(request, 'dashboard/hod_submission_detail.html', context)


def get_status_class(status):
    """Get Bootstrap class for submission status."""
    status_classes = {
        'Draft': 'secondary',
        'Call Issued': 'secondary',
        'HOD/DM Submit': 'primary',
        'Review of Procurement Draft': 'primary',
        'Submit Compiled Document': 'warning',
        'CBM Review': 'primary',
        'Returned': 'warning',
        'Publish Plan': 'success',
        'Prepare Tender Document': 'success',
        'CBM Review TD': 'primary',
        'Publication of TD': 'success',
        'Opening': 'info',
        'Evaluation': 'info',
        'Notify Bidders': 'info',
        'Contract Negotiation': 'info',
        'Contract Drafting': 'info',
        'Legal Review': 'primary',
        'MINIJUST Legal Review': 'primary',
        'Awarded': 'success',
        'Completed': 'success',
        'Rejected': 'danger',
        'Cancelled': 'danger',
    }
    return status_classes.get(status, 'secondary')


@login_required
def hod_procurement_call_detail_view(request, call_id):
    """
    View detail of a procurement call for HOD/DM.
    Shows call information, submissions for that call, and links to create submission.
    """
    user = request.user
    
    if user.role and user.role.name not in ['HOD/DM', 'Admin']:
        return redirect('dashboard')
    
    from apps.procurement.models import ProcurementCall, Submission
    
    try:
        call = ProcurementCall.objects.get(id=call_id)
    except ProcurementCall.DoesNotExist:
        return render(request, 'dashboard/error.html', {'error': 'Procurement call not found'})
    
    # Get submissions for this division and call
    division = user.division
    submissions = Submission.objects.filter(
        call=call,
        division=division,
        is_deleted=False
    ).order_by('-created_at')
    
    # Calculate remaining days
    from django.utils import timezone
    remaining_days = (call.end_date - timezone.now()).days if call.end_date > timezone.now() else 0
    is_deadline_passed = call.end_date <= timezone.now()
    
    # Detect document file type for preview
    doc_file_type = None
    absolute_document_url = None
    doc_preview_html = None
    if call.call_document:
        filename = call.call_document.name.lower()
        if filename.endswith('.pdf'):
            doc_file_type = 'pdf'
        elif filename.endswith('.txt'):
            doc_file_type = 'txt'
        elif filename.endswith(('.doc', '.docx')):
            doc_file_type = 'word'
            # Convert Word to HTML
            try:
                doc_preview_html = convert_word_to_html(call.call_document.path)
            except Exception as e:
                doc_preview_html = f'<div class="alert alert-danger">Error loading Word document preview: {str(e)}</div>'
        elif filename.endswith(('.xls', '.xlsx')):
            doc_file_type = 'excel'
            # Convert Excel to HTML
            try:
                doc_preview_html = convert_excel_to_html(call.call_document.path)
            except Exception as e:
                doc_preview_html = f'<div class="alert alert-danger">Error loading Excel file preview: {str(e)}</div>'
        elif filename.endswith('.zip'):
            doc_file_type = 'zip'
        # Convert relative URL to absolute URL for Google Docs Viewer
        absolute_document_url = request.build_absolute_uri(call.call_document.url)
    
    context = {
        'call': call,
        'submissions': submissions,
        'division': division,
        'remaining_days': remaining_days,
        'is_deadline_passed': is_deadline_passed,
        'doc_file_type': doc_file_type,
        'absolute_document_url': absolute_document_url,
        'doc_preview_html': doc_preview_html,
    }
    
    return render(request, 'dashboard/hod_procurement_call_detail.html', context)


@login_required
def hod_edit_submission_view(request, submission_id):
    """
    Edit an existing submission - simplified form.
    Only draft and returned submissions can be edited.
    """
    from apps.procurement.models import Submission, SubmissionDocument
    
    user = request.user
    
    if user.role and user.role.name not in ['HOD/DM', 'Admin']:
        return redirect('dashboard')
    
    if request.method == 'POST':
        try:
            submission = Submission.objects.get(id=submission_id)
            
            # Check permission - user can only edit their division's submissions
            if user.role and user.role.name == 'HOD/DM' and user.division != submission.division:
                return render(request, 'dashboard/error.html', {'error': 'Access denied'})
            
            # Check if submission is editable
            if submission.status not in ['Draft', 'Returned']:
                return render(request, 'dashboard/error.html', 
                            {'error': f'Cannot edit submission with status: {submission.status}'})
            
            # Get form data
            number_of_items_str = request.POST.get('number_of_procurement_items', '').strip()
            item_description = request.POST.get('item_description', '').strip()
            
            # Validate required fields
            if not number_of_items_str:
                return render(request, 'dashboard/error.html', {'error': 'Number of items is required'})
            if not item_description:
                return render(request, 'dashboard/error.html', {'error': 'Description is required'})
            
            # Validate number of items
            try:
                number_of_items = int(number_of_items_str)
                if number_of_items <= 0:
                    return render(request, 'dashboard/error.html', {'error': 'Number of items must be greater than 0'})
            except (ValueError, TypeError):
                return render(request, 'dashboard/error.html', {'error': 'Invalid number of items'})
            
            # Update submission fields
            submission.item_description = item_description
            submission.number_of_procurement_items = number_of_items
            submission.save()
            
            # Process supporting documents (replace existing with new uploads)
            document_types = {
                'procurement_plan_files': 'procurement_plan',
                'technical_specification_files': 'technical_specification',
                'market_survey_files': 'market_survey',
            }
            
            for form_field, doc_type in document_types.items():
                if form_field in request.FILES and request.FILES.getlist(form_field):
                    # Delete existing documents of this type
                    submission.supporting_documents.filter(document_type=doc_type).delete()
                    
                    # Add new documents
                    for uploaded_file in request.FILES.getlist(form_field):
                        try:
                            # Check file size (max 10MB)
                            if uploaded_file.size > 10 * 1024 * 1024:
                                continue  # Skip file if too large
                            
                            # Create SubmissionDocument record
                            SubmissionDocument.objects.create(
                                submission=submission,
                                document_type=doc_type,
                                file=uploaded_file,
                                original_filename=uploaded_file.name,
                                file_size=uploaded_file.size,
                                uploaded_by=user,
                            )
                        except Exception as file_error:
                            print(f"Error saving document: {file_error}")
                            continue
            
            # Show success message
            from django.contrib import messages
            messages.success(request, f'Submission {submission.tracking_reference} updated successfully!')
            
            return redirect('hod_submission_detail', submission_id=submission.id)
        
        except Submission.DoesNotExist:
            return render(request, 'dashboard/error.html', {'error': 'Submission not found'})
        except Exception as e:
            import traceback
            print(f"Error editing submission: {e}")
            print(traceback.format_exc())
            return render(request, 'dashboard/error.html', {'error': f'An error occurred: {str(e)}'})
    
    # For GET requests, redirect to submission detail
    return redirect('hod_submission_detail', submission_id=submission_id)


@login_required
def hod_create_submission_view(request):
    """
    Create a new submission for a procurement call with supporting documents.
    Simplified form: Call, Division, Number of Items, Description, and 3 document types.
    """
    from apps.procurement.models import Submission, ProcurementCall, SubmissionDocument
    from apps.workflows.models import WorkflowStage
    from decimal import Decimal
    
    user = request.user
    
    if user.role and user.role.name not in ['HOD/DM', 'Admin']:
        return redirect('dashboard')
    
    division = user.division
    if not division:
        return render(request, 'dashboard/error.html', {'error': 'No division assigned'})
    
    if request.method == 'POST':
        try:
            call_id = request.POST.get('call_id', '').strip()
            number_of_items_str = request.POST.get('number_of_procurement_items', '').strip()
            item_description = request.POST.get('item_description', '').strip()
            
            # Validate required fields
            if not call_id:
                return render(request, 'dashboard/error.html', {'error': 'Procurement call is required'})
            if not number_of_items_str:
                return render(request, 'dashboard/error.html', {'error': 'Number of procurement items is required'})
            if not item_description:
                return render(request, 'dashboard/error.html', {'error': 'Description is required'})
            
            # Validate number of items
            try:
                number_of_items = int(number_of_items_str)
                if number_of_items <= 0:
                    return render(request, 'dashboard/error.html', {'error': 'Number of items must be greater than 0'})
            except (ValueError, TypeError):
                return render(request, 'dashboard/error.html', {'error': 'Invalid number of items'})
            
            # Get the procurement call
            try:
                call = ProcurementCall.objects.get(id=call_id)
            except ProcurementCall.DoesNotExist:
                return render(request, 'dashboard/error.html', {'error': 'Procurement call not found'})
            
            # Create submission with simplified fields
            submission = Submission.objects.create(
                call=call,
                division=division,
                item_name=f"{division.name} - {call.title[:50]}",  # Auto-generated item name
                item_description=item_description,
                number_of_procurement_items=number_of_items,
                quantity=1,  # Default quantity
                unit_of_measure='Items',  # Default unit
                estimated_unit_price=Decimal('0.00'),  # Will be set during budget planning
                submitted_by=user,
                status='Draft',
            )
            
            # Set initial workflow stage
            initial_stage = WorkflowStage.objects.filter(order=1).first()
            if initial_stage:
                submission.current_stage = initial_stage
                submission.save()
            
            # Process supporting documents
            document_types = {
                'procurement_plan_files': 'procurement_plan',
                'technical_specification_files': 'technical_specification',
                'market_survey_files': 'market_survey',
            }
            
            for form_field, doc_type in document_types.items():
                if form_field in request.FILES:
                    for uploaded_file in request.FILES.getlist(form_field):
                        try:
                            # Check file size (max 10MB)
                            if uploaded_file.size > 10 * 1024 * 1024:
                                continue  # Skip file if too large
                            
                            # Create SubmissionDocument record
                            doc = SubmissionDocument.objects.create(
                                submission=submission,
                                document_type=doc_type,
                                file=uploaded_file,
                                original_filename=uploaded_file.name,
                                file_size=uploaded_file.size,
                                uploaded_by=user,
                            )
                        except Exception as file_error:
                            print(f"Error saving document: {file_error}")
                            continue
            
            # Show success message
            from django.contrib import messages
            messages.success(request, f'Submission {submission.tracking_reference} created successfully!')
            
            return redirect('hod_submission_detail', submission_id=submission.id)
        
        except Exception as e:
            import traceback
            print(f"Error creating submission: {e}")
            print(traceback.format_exc())
            return render(request, 'dashboard/error.html', {'error': f'An error occurred: {str(e)}'})
    
    return redirect('hod_submissions')


@login_required
def hod_submit_submission_view(request, submission_id):
    """
    Submit a draft submission for Procurement Team review.
    Moves submission from Draft to HOD/DM Submit (stage 2), then to Review of Procurement Draft (stage 3).
    """
    from apps.procurement.models import Submission
    from apps.workflows.models import WorkflowStage

    user = request.user
    
    if user.role and user.role.name not in ['HOD/DM', 'Admin']:
        return redirect('dashboard')
    
    try:
        submission = Submission.objects.get(id=submission_id)
        
        # Check permission
        if user.role and user.role.name == 'HOD/DM' and user.division != submission.division:
            return render(request, 'dashboard/error.html', {'error': 'Access denied'})
        
        # Only allow submission from Draft or Returned status
        if submission.status not in ['Draft', 'Returned']:
            return render(request, 'dashboard/error.html', 
                        {'error': f'Cannot submit submission with status: {submission.status}'})
        
        # Update submission status to HOD/DM Submit but move to Review of Procurement Draft stage
        # Status shows who submitted, Stage shows who should review next
        submission.status = 'HOD/DM Submit'
        submission.submitted_at = timezone.now()
        submission.submitted_by = user
        
        # Set current stage to "Review of Procurement Draft" (stage 3) - waiting for Procurement Team
        review_draft_stage = WorkflowStage.objects.filter(order=3).first()
        if review_draft_stage:
            submission.current_stage = review_draft_stage
        
        submission.save()
        
        # Log the action
        from apps.workflows.models import WorkflowHistory
        previous_stage = None
        if submission.current_stage and submission.current_stage.order > 1:
            previous_stage = WorkflowStage.objects.filter(order=submission.current_stage.order - 1).first()
        
        WorkflowHistory.objects.create(
            submission=submission,
            from_stage=previous_stage,
            to_stage=submission.current_stage,
            action='submit',
            comments='Division submitted procurement request for Procurement Team review',
            action_by=user,
        )
        
        # Create notifications for Procurement Team
        from apps.notifications.models import Notification
        from apps.accounts.models import User
        
        procurement_users = User.objects.filter(role__name='Procurement Team', is_active=True)
        
        for reviewer in procurement_users:
            Notification.objects.create(
                user=reviewer,
                title=f'New Submission for Review: {submission.tracking_reference}',
                message=f'{submission.division.name} submitted a procurement request for {submission.item_name}. Please review the draft.',
                notification_type='submission_status',
                priority='high',
                related_object_type='Submission',
                related_object_id=str(submission.id),
                action_url=f'/dashboard/procurement/submissions/{submission.id}/',
            )
        
        # Show success message
        from django.contrib import messages
        messages.success(request, f'Submission {submission.tracking_reference} submitted for Procurement Team review!')
        
        return redirect('hod_submission_detail', submission_id=submission.id)
    
    except Submission.DoesNotExist:
        return render(request, 'dashboard/error.html', {'error': 'Submission not found'})
    except Exception as e:
        return render(request, 'dashboard/error.html', {'error': str(e)})


@login_required
def hod_submit_clarification_view(request, submission_id):
    """
    Submit clarification/response to CBM feedback.
    """
    from apps.procurement.models import Submission
    from apps.workflows.models import WorkflowHistory
    
    user = request.user
    
    if user.role and user.role.name not in ['HOD/DM', 'Admin']:
        return redirect('dashboard')
    
    if request.method == 'POST':
        try:
            submission = Submission.objects.get(id=submission_id)
            
            # Check permission
            if user.role and user.role.name == 'HOD/DM' and user.division != submission.division:
                return render(request, 'dashboard/error.html', {'error': 'Access denied'})
            
            # Only allow for Returned status
            if submission.status != 'Returned':
                return render(request, 'dashboard/error.html', 
                            {'error': 'This submission is not awaiting clarification'})
            
            # Get the response
            response_text = request.POST.get('clarification_response', '')
            
            if not response_text:
                return render(request, 'dashboard/error.html', 
                            {'error': 'Please provide a response'})
            
            # Create comment with the response
            from apps.procurement.models import Comment
            Comment.objects.create(
                submission=submission,
                author=user,
                content=response_text,
            )
            
            # Handle file attachments if any
            if 'clarification_attachment' in request.FILES:
                import os
                import uuid
                for uploaded_file in request.FILES.getlist('clarification_attachment'):
                    file_name = f"submissions/{submission.id}/{uuid.uuid4()}_{uploaded_file.name}"
                    file_path = os.path.join(settings.MEDIA_ROOT, file_name)
                    os.makedirs(os.path.dirname(file_path), exist_ok=True)
                    
                    with open(file_path, 'wb+') as destination:
                        for chunk in uploaded_file.chunks():
                            destination.write(chunk)
                    
                    attachment = {
                        'name': uploaded_file.name,
                        'url': f'{settings.MEDIA_URL}{file_name}'
                    }
                    if submission.attachments:
                        submission.attachments.append(attachment)
                    else:
                        submission.attachments = [attachment]
            
            # Update submission status back to Submitted for re-review
            submission.status = 'Submitted'
            submission.save()
            
            # Log the action
            WorkflowHistory.objects.create(
                submission=submission,
                from_stage=submission.current_stage,
                to_stage=submission.current_stage,
                action='clarify',
                comments='Division provided clarifications',
                action_by=user,
            )
            
            # Notify the appropriate reviewer based on current stage
            from apps.notifications.models import Notification
            from apps.accounts.models import User
            
            # Determine who should be notified based on the current stage
            notification_users = []
            
            if submission.current_stage:
                if submission.current_stage.order == 3:  # Review of Procurement Draft
                    notification_users = list(User.objects.filter(role__name='Procurement Team', is_active=True))
                elif submission.current_stage.order == 4:  # CBM Review
                    notification_users = list(User.objects.filter(role__name='CBM', is_active=True))
                elif submission.current_stage.order == 7:  # CBM Review TD
                    notification_users = list(User.objects.filter(role__name='CBM', is_active=True))
                elif submission.current_stage.order in [14, 15]:  # Legal Reviews
                    notification_users = list(User.objects.filter(role__name='CBM', is_active=True)) + \
                                        list(User.objects.filter(role__name='Admin', is_active=True))
            
            # If we couldn't determine, notify Procurement Team (most common reviewer)
            if not notification_users:
                notification_users = list(User.objects.filter(role__name='Procurement Team', is_active=True))
            
            for reviewer in notification_users:
                Notification.objects.create(
                    user=reviewer,
                    title=f'Clarification Provided: {submission.tracking_reference}',
                    message=f'{submission.division.name} has provided clarifications for {submission.item_name}',
                    notification_type='submission_status',
                    priority='medium',
                    related_object_type='Submission',
                    related_object_id=str(submission.id),
                    action_url=f'/dashboard/cbm/submissions/{submission.id}/',
                )
            
            from django.contrib import messages
            messages.success(request, 'Clarifications submitted successfully!')
            
            return redirect('hod_submission_detail', submission_id=submission.id)
        
        except Submission.DoesNotExist:
            return render(request, 'dashboard/error.html', {'error': 'Submission not found'})
        except Exception as e:
            return render(request, 'dashboard/error.html', {'error': str(e)})
    
    return redirect('hod_submission_detail', submission_id=submission_id)
